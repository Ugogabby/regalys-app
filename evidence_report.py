"""
evidence_report.py
───────────────────
REgalys Evidence Report Generator

Packages a completed REgalys query session into a downloadable evidence
report artifact. Outputs:

  1. In-memory bytes → Streamlit st.download_button (Word .docx)

The report includes:
  - Query and run metadata
  - Full synthesis answer (markdown → formatted Word paragraphs)
  - HTA-style evidence table (PICO + quality flags per paper)
  - Quality summary statistics
  - Full reference list with PMIDs

This is the direct deliverable a HEOR scientist would attach to an
HTA dossier or include in a systematic literature review.

Usage:
    from evidence_report import build_word_report
    docx_bytes = build_word_report(query, answer, evidence_table, quality_summary)

    # In Streamlit:
    st.download_button("📥 Download Evidence Report", docx_bytes,
                       file_name="regalys_evidence_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
"""

import io
import re
from datetime import datetime
from typing import Optional


# ─────────────────────────────────────────────────────────────────────────────
# Word document builder (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _try_import_docx():
    """Lazy import — only required if report download is used."""
    try:
        import docx
        return docx
    except ImportError:
        return None


def _add_heading(doc, text: str, level: int = 1):
    """Add a heading with the correct level."""
    doc.add_heading(text, level=level)


def _add_paragraph(doc, text: str, bold: bool = False, style: str = None):
    """Add a paragraph, optionally bold."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def _markdown_to_paragraphs(doc, markdown_text: str):
    """
    Converts the REgalys synthesis answer (markdown) into Word paragraphs.
    Handles: ## headings, bullet points, bold text, plain paragraphs.
    Tables are skipped (too complex for this lightweight converter).
    """
    lines = markdown_text.split("\n")

    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        # Heading detection
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # Bullet / numbered list detection
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            # Strip inline bold markers
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")

        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")

        # Table rows — skip (render as plain text instead)
        elif line.startswith("|"):
            # Strip pipes and render as monospace-ish paragraph
            cells = [c.strip() for c in line.strip("|").split("|")]
            doc.add_paragraph("  |  ".join(cells))

        else:
            # Regular paragraph — strip inline bold
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            doc.add_paragraph(text)


def _build_evidence_table_in_doc(doc, evidence_table: list[dict]):
    """
    Inserts the HTA evidence table as a formatted Word table.

    Columns are ordered to match standard HEOR dossier format:
    Study | Year | Design | Data Source | P | I | C | O | Quality flags
    """
    if not evidence_table:
        doc.add_paragraph("No evidence table rows met the confidence threshold.")
        return

    # Define column order — show most important columns first
    priority_cols = [
        "Title", "Year", "PMID", "Study Design", "Data Source",
        "Population", "Intervention", "Comparator", "Outcome",
        "Time Horizon", "PICO Confidence",
        "Active Comparator", "New User Design", "Validated Outcome",
        "Confounding Method", "Competing Events", "Immortal Time Protected",
        "Overall Quality",
    ]

    # Determine which columns actually have data
    all_keys = set()
    for row in evidence_table:
        all_keys.update(row.keys())
    columns = [c for c in priority_cols if c in all_keys]

    # Create table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        # Bold header
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for row_data in evidence_table:
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row_data.get(col, "")
            row_cells[i].text = str(val) if val else "—"


def _build_quality_summary_section(doc, quality_summary: dict):
    """
    Inserts the quality summary statistics as a bulleted list.
    """
    if not quality_summary or quality_summary.get("total_assessed", 0) == 0:
        doc.add_paragraph("No quality assessment data available.")
        return

    n = quality_summary["total_assessed"]
    doc.add_paragraph(f"Papers assessed: {n}", style="List Bullet")
    doc.add_paragraph(f"Active comparator design: {quality_summary.get('active_comparator', '—')}", style="List Bullet")
    doc.add_paragraph(f"New user design: {quality_summary.get('new_user_design', '—')}", style="List Bullet")
    doc.add_paragraph(f"Validated outcome: {quality_summary.get('validated_outcome', '—')}", style="List Bullet")
    doc.add_paragraph(f"Competing events addressed: {quality_summary.get('competing_events_handled', '—')}", style="List Bullet")
    doc.add_paragraph(f"Immortal time protected: {quality_summary.get('immortal_time_protected', '—')}", style="List Bullet")

    dist = quality_summary.get("quality_distribution", {})
    if dist:
        doc.add_paragraph(
            f"Quality distribution — Strong: {dist.get('strong',0)}, "
            f"Moderate: {dist.get('moderate',0)}, Weak: {dist.get('weak',0)}",
            style="List Bullet"
        )


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def build_word_report(
    query:           str,
    answer:          str,
    evidence_table:  list[dict],
    quality_summary: dict,
    chunks:          list[dict] = None,
    timing:          dict       = None,
) -> Optional[bytes]:
    """
    Builds a formatted Word document evidence report.

    Args:
        query:           the original user query
        answer:          the Claude synthesis answer (markdown format)
        evidence_table:  list of evidence table row dicts from PICOEnricher
        quality_summary: quality statistics dict from PICOEnricher
        chunks:          original retrieved chunks (for reference list)
        timing:          timing dict for metadata

    Returns:
        Word document as bytes, or None if python-docx is not installed.
    """
    docx = _try_import_docx()
    if docx is None:
        print("  [EvidenceReport] python-docx not installed — cannot generate report")
        return None

    doc = docx.Document()

    # ── Title page ─────────────────────────────────────────────────────────────
    _add_heading(doc, "REgalys Evidence Report", level=1)
    _add_heading(doc, "Real-world Evidence Generation and Analysis Insights", level=2)

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}"
    )
    doc.add_paragraph(
        "Built by Ugochukwu Ezigbo BPharm, MHA | "
        "PhD Candidate, Pharmaceutical Outcomes & Policy | "
        "University of Pittsburgh"
    )
    doc.add_paragraph(
        "Knowledge base: 3,894 pharmacoepidemiology papers · 302,377 chunks · "
        "Voyage AI embeddings · Pinecone · Cohere Rerank 3 · Claude Sonnet"
    )

    doc.add_page_break()

    # ── Query ─────────────────────────────────────────────────────────────────
    _add_heading(doc, "1. Query", level=1)
    doc.add_paragraph(query)

    if timing:
        total = sum(timing.get(k, 0) for k in ["retrieval", "rerank", "llm"])
        doc.add_paragraph(
            f"Retrieval: {timing.get('retrieval', 0):.1f}s · "
            f"Rerank: {timing.get('rerank', 0):.1f}s · "
            f"Synthesis: {timing.get('llm', 0):.1f}s · "
            f"Total: {total:.1f}s · "
            f"Chunks used: {timing.get('chunks', 0)}"
        )

    # ── Synthesis answer ──────────────────────────────────────────────────────
    _add_heading(doc, "2. Synthesis", level=1)
    _markdown_to_paragraphs(doc, answer)

    # ── Evidence table ────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "3. Evidence Table", level=1)
    doc.add_paragraph(
        "The following table summarizes PICO elements and study quality flags "
        "extracted from retrieved papers. Quality flags use pharmacoepidemiology "
        "standards: active comparator new-user design, validated outcome definitions, "
        "and appropriate confounding control."
    )
    _build_evidence_table_in_doc(doc, evidence_table)

    # ── Quality summary ───────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "4. Study Quality Summary", level=1)
    doc.add_paragraph(
        "Aggregate quality statistics across all retrieved papers assessed "
        "for this query."
    )
    _build_quality_summary_section(doc, quality_summary)

    # ── Reference list ────────────────────────────────────────────────────────
    if chunks:
        doc.add_page_break()
        _add_heading(doc, "5. Retrieved Sources", level=1)
        doc.add_paragraph(
            f"{len(chunks)} chunks retrieved from the REgalys knowledge base."
        )
        seen_pmids = set()
        ref_num = 1
        for chunk in chunks:
            pmid     = str(chunk.get("pmid",   ""))
            citation = chunk.get("citation",   "")
            title    = chunk.get("title",      "")[:120]
            authors  = chunk.get("authors",    "")
            year     = chunk.get("year",       "")
            journal  = chunk.get("journal",    "")

            # Deduplicate by PMID
            if pmid and pmid in seen_pmids:
                continue
            if pmid:
                seen_pmids.add(pmid)

            ref_text = citation or f"{authors} ({year}). {title}. {journal}."
            if pmid and not pmid.startswith("BOOK_"):
                ref_text += f" PMID: {pmid}"

            doc.add_paragraph(f"[{ref_num}] {ref_text}", style="List Number")
            ref_num += 1

    # ── Disclaimer ────────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "Disclaimer", level=2)
    doc.add_paragraph(
        "This report was generated by REgalys, an AI-powered evidence synthesis "
        "system for pharmacoepidemiology research. The synthesis is based on "
        "retrieved literature from the REgalys knowledge base and should be "
        "reviewed by a qualified researcher before use in regulatory submissions, "
        "HTA dossiers, or clinical decision-making. PICO and quality assessments "
        "are automated extractions and may contain errors — manual verification "
        "of critical claims is recommended."
    )

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
